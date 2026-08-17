#!/usr/bin/env python3
"""
Document Generator for Explorer-d334
Generates .txt, .docx, .md files from content or AI generation
"""

import os
import json
from pathlib import Path
from datetime import datetime
import hashlib

# Try to import optional dependencies
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import markdown
    MD_AVAILABLE = True
except ImportError:
    MD_AVAILABLE = False

class DocumentGenerator:
    def __init__(self):
        self.output_dir = Path("generated_docs")
        self.output_dir.mkdir(exist_ok=True)
    
    def generate_txt(self, filename, content):
        """Generate a plain text file"""
        filepath = self.output_dir / f"{filename}.txt"
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)
        return filepath
    
    def generate_md(self, filename, content):
        """Generate a markdown file"""
        filepath = self.output_dir / f"{filename}.md"
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)
        
        # Optionally convert to HTML
        if MD_AVAILABLE:
            html_content = markdown.markdown(content)
            html_path = self.output_dir / f"{filename}.html"
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(f"<!DOCTYPE html><html><head><meta charset='utf-8'><title>{filename}</title></head><body>{html_content}</body></html>")
        
        return filepath
    
    def generate_docx(self, filename, content, title=None):
        """Generate a Word document"""
        if not DOCX_AVAILABLE:
            return None
        
        doc = Document()
        if title:
            doc.add_heading(title, 0)
        
        # Split content into paragraphs
        for paragraph in content.split('\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph)
        
        filepath = self.output_dir / f"{filename}.docx"
        doc.save(filepath)
        return filepath
    
    def generate_from_ai(self, topic, doc_type='md'):
        """Generate document using AI (Gemma)"""
        try:
            import subprocess
            prompt = f"Write a comprehensive document about {topic}. Include an introduction, main sections, and conclusion. Format it nicely with markdown headers."
            
            result = subprocess.run(
                ["ollama", "run", "gemma2:2b", prompt],
                capture_output=True,
                text=True,
                timeout=60
            )
            
            if result.returncode == 0 and result.stdout.strip():
                content = result.stdout.strip()
                
                # Save to SCP memory
                self.save_to_memory(topic, content, doc_type)
                
                # Generate the document
                if doc_type == 'txt':
                    filepath = self.generate_txt(topic, content)
                elif doc_type == 'md':
                    filepath = self.generate_md(topic, content)
                elif doc_type == 'docx':
                    filepath = self.generate_docx(topic, content, topic)
                else:
                    filepath = self.generate_txt(topic, content)
                
                return {
                    'success': True,
                    'topic': topic,
                    'type': doc_type,
                    'filepath': str(filepath),
                    'word_count': len(content.split())
                }
        except Exception as e:
            return {'success': False, 'error': str(e)}
        
        return {'success': False, 'error': 'AI generation failed'}
    
    def generate_from_template(self, template_name, data, doc_type='md'):
        """Generate document from template"""
        templates = {
            'report': "# Report: {title}\n\n## Date\n{date}\n\n## Summary\n{summary}\n\n## Details\n{details}\n\n## Conclusion\n{conclusion}",
            'note': "# Note: {title}\n\n{content}\n\n---\n*Created: {date}*",
            'article': "# {title}\n\n## Introduction\n{introduction}\n\n## Body\n{body}\n\n## Conclusion\n{conclusion}\n\n---\n*By Explorer-d334*",
            'readme': "# {project_name}\n\n## Description\n{description}\n\n## Installation\n{installation}\n\n## Usage\n{usage}\n\n## License\n{license}"
        }
        
        template = templates.get(template_name, templates['note'])
        
        # Add date if not present
        if '{date}' in template and 'date' not in data:
            data['date'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # Fill template
        content = template.format(**data)
        
        # Generate filename from title
        title = data.get('title', template_name)
        filename = title.lower().replace(' ', '_').replace('/', '_')[:50]
        
        if doc_type == 'txt':
            filepath = self.generate_txt(filename, content)
        elif doc_type == 'md':
            filepath = self.generate_md(filename, content)
        elif doc_type == 'docx':
            filepath = self.generate_docx(filename, content, title)
        else:
            filepath = self.generate_txt(filename, content)
        
        # Save to memory
        self.save_to_memory(filename, content, doc_type)
        
        return {
            'success': True,
            'filename': filename,
            'type': doc_type,
            'filepath': str(filepath),
            'template': template_name
        }
    
    def save_to_memory(self, title, content, doc_type):
        """Save document info to SCP memory"""
        try:
            from scp_memory import get_scp_memory
            memory = get_scp_memory()
            memory.create_scp("generated_document", title, {
                "type": doc_type,
                "content": content[:500],
                "generated_at": datetime.now().isoformat()
            })
        except:
            pass
    
    def list_documents(self):
        """List all generated documents"""
        docs = []
        for ext in ['*.txt', '*.md', '*.docx', '*.html']:
            for file in self.output_dir.glob(ext):
                docs.append({
                    'name': file.name,
                    'size': file.stat().st_size,
                    'modified': datetime.fromtimestamp(file.stat().st_mtime).isoformat()[:19]
                })
        return docs

if __name__ == "__main__":
    import sys
    
    gen = DocumentGenerator()
    
    if len(sys.argv) < 2:
        print("Document Generator Commands:")
        print("  ai <topic> [type]      - Generate AI document")
        print("  template <name> <json> - Generate from template")
        print("  txt <file> <content>   - Generate text file")
        print("  md <file> <content>    - Generate markdown")
        print("  list                   - List generated docs")
    
    elif sys.argv[1] == "ai":
        topic = sys.argv[2]
        doc_type = sys.argv[3] if len(sys.argv) > 3 else 'md'
        result = gen.generate_from_ai(topic, doc_type)
        if result.get('success'):
            print(f"✅ Generated {result['type']}: {result['filepath']}")
            print(f"   Words: {result['word_count']}")
        else:
            print(f"❌ Failed: {result.get('error')}")
    
    elif sys.argv[1] == "template":
        template = sys.argv[2]
        data = json.loads(sys.argv[3]) if len(sys.argv) > 3 else {}
        result = gen.generate_from_template(template, data)
        if result.get('success'):
            print(f"✅ Generated: {result['filepath']}")
    
    elif sys.argv[1] == "list":
        docs = gen.list_documents()
        for doc in docs:
            print(f"  📄 {doc['name']} ({doc['size']} bytes)")
    
    gen = DocumentGenerator()
